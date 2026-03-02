import sys
import os
import subprocess

def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

install_and_import('docx')
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)
    
    # A workaround for the lack of a hyperlink style
    r.font.color.theme_color = docx.enum.dml.MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def create_report():
    doc = Document()
    
    # Styling for heading
    h1 = doc.add_heading('Laporan Tugas Kriptografi: Kalkulator Enkripsi-Dekripsi Berbasis Web', 0)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Nama: Darren Nathanael Melakha')
    doc.add_paragraph('NIM: 21120123120001')
    doc.add_paragraph('Kelas: Kriptografi C')
    doc.add_paragraph('Bahasa Pemrograman: Javascript, HTML, CSS\n')

    # Section 1: Source Code
    doc.add_heading('1. Source Program', level=1)
    
    files = {
        'index.html': 'Antarmuka Pengguna Utama (HTML)',
        'style.css': 'Desain Tampilan (CSS)',
        'script.js': 'Logika Enkripsi & Dekripsi 5 Cipher (Javascript)'
    }
    
    for filename, description in files.items():
        doc.add_heading(f'File: {filename} - {description}', level=2)
        try:
            with open(filename, 'r', encoding='utf-8') as f:
                content = f.read()
                # Adding code inside a separate font format (simulating monospace)
                p = doc.add_paragraph()
                run = p.add_run(content)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
        except Exception as e:
            doc.add_paragraph(f'[Error membaca file {filename}: {e}]')

    # Section 2: Screenshots Placeholder
    doc.add_heading('2. Tampilan Antarmuka Program', level=1)
    doc.add_paragraph('Silakan ' + 'insert' + '/paste hasil print screen antarmuka (GUI) program di bawah garis ini.')
    doc.add_paragraph('--------------------------------------------------')
    doc.add_paragraph('[ Ganti baris ini dengan Gambar Print Screen Interface ]')
    doc.add_paragraph('--------------------------------------------------\n')
    
    # Section 3: Example Plaintext/Ciphertext
    doc.add_heading('3. Contoh Plainteks dan Cipherteks', level=1)
    doc.add_paragraph('Berikut adalah contoh proses pengujian aplikasi menggunakan tipe data teks.')
    
    examples = [
        ("Vigenere Cipher", "HELLO DARREN", "KEY", "RIJVS HKVVBX"),
        ("Affine Cipher", "KRIPTOGRAFI C", "(a=5, b=8)", "KVSDFUKVOPU I"),
        ("Playfair Cipher", "KRIPTOGRAFI C", "KEYWORD", "OCNKWUKCPNCD R"),
        ("Hill Cipher", "DARREN", "[3 3; 2 5]", "OHZHYX"),
        ("Enigma Cipher", "KRIPTOGRAFI C", "Rotor I(A), II(A), III(A)", "LVDWGHMOWKJ V")
    ]
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Jenis Cipher'
    hdr_cells[1].text = 'Plainteks'
    hdr_cells[2].text = 'Key / Konfigurasi'
    hdr_cells[3].text = 'Hasil Cipherteks'
    
    for cipher, text, key, res in examples:
        row_cells = table.add_row().cells
        row_cells[0].text = cipher
        row_cells[1].text = text
        row_cells[2].text = key
        row_cells[3].text = res
        
    doc.add_paragraph('\n[Catatan untuk Darren: Anda juga dapat menambahkan ilustrasi print screen contoh hasil enkripsi gambar, database, audio, atau video sesuai permintaan dosen di bagian ini]\n')

    # Section 4: Repository Link Placeholder
    doc.add_heading('4. Link Repository Github / Google Drive', level=1)
    p = doc.add_paragraph('Repository Kode Program: ')
    # add_hyperlink(p, 'Link ke Repository (Ganti dengan link Anda)', 'https://github.com/darren-nm/Tugas_Kripto')
    p.add_run('[Ganti dengan Link Github atau Link Google Drive Anda here]')

    # Save Document
    doc.save('Laporan_Kriptografi_Darren.docx')
    print("Document saved successfully!")

if __name__ == '__main__':
    create_report()
